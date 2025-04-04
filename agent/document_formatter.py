# agent/document_formatter.py
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor # Import measurement classes
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT # Import enums
from docx.enum.section import WD_SECTION_START, WD_HEADER_FOOTER # For page numbering/breaks
from docx.enum.style import WD_STYLE_TYPE # If using built-in styles
from docx.oxml.ns import qn # For XML querying if needed
from docx.oxml import OxmlElement # For direct XML manipulation (used for page numbers)

from pathlib import Path
import os # Needed for checking paths

# Assuming GuidelineManager is importable
from .guideline_manager import GuidelineManager

# --- Placeholder Constants ---
TOC_PLACEHOLDER = "[---TABLE_OF_CONTENTS---]"
LOF_PLACEHOLDER = "[---LIST_OF_FIGURES---]"
LOT_PLACEHOLDER = "[---LIST_OF_TABLES---]"

# Default Color Constant (Example)
COLOR_BLACK = RGBColor(0, 0, 0)

class DocumentFormatter:
    """
    Handles the creation, formatting, and finalization of the .docx document
    using python-docx, based on rules provided by GuidelineManager.

    Includes title page, declaration, auto-numbered headings/figures/tables,
    placeholders, dynamic list generation (TOC, LoF, LoT with placeholders),
    and section-based page numbering (Roman/Arabic).
    """
    def __init__(self, guideline_manager: GuidelineManager):
        self.guideline_mgr = guideline_manager
        self.doc = None
        self.current_section = None
        self.headings = []
        self.figures = []
        self.tables = []
        self.chapter_num = 0
        self.section_num = 0
        self.subsection_num = 0
        self.figure_num_in_chapter = 0
        self.table_num_in_chapter = 0
        self.current_chapter_number = 0
        self.placeholder_paragraphs = {}
        self.front_matter_section_index = 0
        self.body_section_index = -1

    def _reset_tracking(self):
        print("    Resetting document tracking lists and counters.")
        self.headings = []
        self.figures = []
        self.tables = []
        self.chapter_num = 0
        self.section_num = 0
        self.subsection_num = 0
        self.figure_num_in_chapter = 0
        self.table_num_in_chapter = 0
        self.current_chapter_number = 0
        self.placeholder_paragraphs = {}
        self.front_matter_section_index = 0
        self.body_section_index = -1

    def create_document(self, doc_type: str):
        self._reset_tracking()
        self.doc = Document()
        self.current_section = self.doc.sections[0]
        self.front_matter_section_index = 0
        self.current_section.page_width = Cm(21.0)
        self.current_section.page_height = Cm(29.7)
        print(f"    Document created. Page size set to A4. Tracking reset.")
        self.apply_margins(doc_type)
        return self.doc

    def apply_margins(self, doc_type: str):
        margins = self.guideline_mgr.get_margins(doc_type)
        if not margins:
            print(f"    Warning: Margin rules not found for {doc_type}. Using defaults.")
            return
        try:
            section = self.doc.sections[-1]
            section.top_margin = margins.get('top', Inches(1.0))
            section.bottom_margin = margins.get('bottom', Inches(1.0))
            section.left_margin = margins.get('left', Inches(1.25))
            section.right_margin = margins.get('right', Inches(1.0))
            print(f"    Margins applied to Section {len(self.doc.sections)-1} for {doc_type}: T={section.top_margin.inches:.2f}\", B={section.bottom_margin.inches:.2f}\", L={section.left_margin.inches:.2f}\", R={section.right_margin.inches:.2f}\"")
        except Exception as e: print(f"    Error applying margins: {e}")

    def _apply_paragraph_format(self, paragraph, style_key: str, doc_type: str):
        style_rules = self.guideline_mgr.get_formatting_rule(doc_type, style_key)
        if not style_rules:
             print(f"    Warning: Style rule '{style_key}' not found for '{doc_type}'. Applying basic format.")
             p_format = paragraph.paragraph_format
             p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
             p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
             p_format.line_spacing = 1.0; p_format.space_before = Pt(0); p_format.space_after = Pt(6)
             for run in paragraph.runs:
                 run.font.name = self.guideline_mgr.get_doc_rules(doc_type).get('font_default', 'Calibri')
                 run.font.size = Pt(11)
             return

        p_format = paragraph.paragraph_format
        if 'align' in style_rules: p_format.alignment = style_rules['align']
        if 'line_spacing' in style_rules:
            p_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p_format.line_spacing = style_rules['line_spacing']
        p_format.space_before = style_rules.get('space_before', Pt(0))
        p_format.space_after = style_rules.get('space_after', Pt(0))
        p_format.first_line_indent = style_rules.get('first_line_indent', None)
        p_format.left_indent = style_rules.get('left_indent', None)
        p_format.right_indent = style_rules.get('right_indent', None)
        p_format.hanging_indent = style_rules.get('hanging_indent', None)
        p_format.keep_together = style_rules.get('keep_together', False)
        p_format.keep_with_next = style_rules.get('keep_with_next', False)
        p_format.page_break_before = style_rules.get('page_break_before', False)
        p_format.widow_control = style_rules.get('widow_control', True)

        base_font_name = style_rules.get('font', self.guideline_mgr.get_doc_rules(doc_type).get('font_default', 'Calibri'))
        base_font_size = style_rules.get('size', Pt(12))
        is_bold = style_rules.get('bold', False); is_italic = style_rules.get('italic', False)
        is_underline = style_rules.get('underline', False); is_all_caps = style_rules.get('all_caps', False)
        font_color = style_rules.get('color', COLOR_BLACK)

        for run in paragraph.runs:
            font = run.font
            font.name = base_font_name; font.size = base_font_size
            font.bold = is_bold; font.italic = is_italic
            font.underline = is_underline; font.all_caps = is_all_caps
            font.color.rgb = font_color

    def add_formatted_paragraph(self, text: str, style_key: str, doc_type: str):
        if text is None: text = ""
        p = self.doc.add_paragraph(str(text))
        self._apply_paragraph_format(p, style_key, doc_type)
        return p

    def add_page_break(self):
        self.doc.add_page_break()

    def add_section_break(self, break_type=WD_SECTION_START.NEW_PAGE):
        self.doc.add_section(break_type)
        self.current_section = self.doc.sections[-1]
        print(f"    Added Section Break. Document now has {len(self.doc.sections)} sections.")
        self.apply_margins('report') # Re-apply margins, assume report context for breaks
        if self.body_section_index == -1 and len(self.doc.sections) > 1:
             self.body_section_index = len(self.doc.sections) - 1
             print(f"    Body Section index set to: {self.body_section_index}")

    def add_title_page(self, doc_type: str, project_data: dict):
        print(f"    Adding Title Page ({doc_type})...")
        layout = self.guideline_mgr.get_title_page_layout(doc_type)
        if not layout: print("Warning: Title page layout not found."); return

        # Simplified Logo Handling
        logo_item = next((item for item in layout if item.get("type") == "logo"), None)
        if logo_item:
            logo_key = self.guideline_mgr._rules.get("common", {}).get("logo_path_key", "logo_image_path")
            logo_path_str = project_data.get(logo_key)
            if logo_path_str:
                logo_path = Path(logo_path_str);
                if not logo_path.is_absolute(): logo_path = (Path.cwd() / logo_path).resolve()
                if logo_path.is_file():
                    try:
                        self.doc.add_picture(str(logo_path), width=Inches(1.5))
                        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self.add_formatted_paragraph("", "title_info", doc_type).paragraph_format.space_after = Pt(18)
                        print(f"      Added logo from: {logo_path}")
                    except Exception as e: print(f"      Warning: Could not add logo: {e}")
                else: print(f"      Warning: Logo file not found: {logo_path}")
            elif not logo_item.get("optional", False): print(f"Warning: Logo required but path missing.")

        # Add Text Elements
        for item in layout:
            if item.get("type") == "logo": continue
            style_key = item.get("style", "normal_text")
            text_to_add = item.get("text", "")
            if not text_to_add and "key" in item:
                data_key = item["key"]
                # Use correct supervisor designation key
                if data_key == 'supervisor_designation':
                    text_to_add = str(project_data.get('supervisor_designation', '[Designation]'))
                else:
                    text_to_add = str(project_data.get(data_key, f"[{data_key}]"))

            final_text = f"{item.get('prefix', '')}{text_to_add}{item.get('suffix', '')}"
            p = self.add_formatted_paragraph(final_text, style_key, doc_type)
            if p:
                if "space_before" in item: p.paragraph_format.space_before = item["space_before"]
                if "space_after" in item: p.paragraph_format.space_after = item["space_after"]

        self.add_page_break(); print("      Title Page added.")

    def add_declaration(self, project_data: dict):
        print("    Adding Declaration Page..."); doc_type = "report"
        template = self.guideline_mgr.get_declaration_text_template()
        if not template or template == "Declaration text not found.": print("Warning: Declaration template not found."); return
        self.add_formatted_paragraph("DECLARATION", "declaration_heading", doc_type)
        try:
            format_data = {k: project_data.get(k, f'[{k}]') for k in ['project_title', 'supervisor_name', 'student_name', 'roll_number']}
            format_data['submission_date'] = project_data.get('submission_month_year', '[Date]')
            declaration_body_text = template.format(**format_data)
            self.add_formatted_paragraph(declaration_body_text, "declaration_body", doc_type)
        except Exception as e: print(f"Error formatting declaration: {e}"); self.add_formatted_paragraph(template, "declaration_body", doc_type)
        self.add_page_break(); print("      Declaration Page added.")

    def add_acknowledgement(self, text: str, doc_type="report"):
        print("    Adding Acknowledgement Page...")
        self.add_formatted_paragraph("ACKNOWLEDGEMENT", "heading_list_toc", doc_type)
        self.add_formatted_paragraph(text or "[Acknowledgement text not generated]", "acknowledgement", doc_type)
        self.add_page_break(); print("      Acknowledgement Page added.")

    def add_abstract(self, text: str, doc_type="report"):
        print("    Adding Abstract Page...")
        self.add_formatted_paragraph("ABSTRACT", "heading_list_toc", doc_type)
        self.add_formatted_paragraph(text or "[Abstract text not generated]", "abstract", doc_type)
        self.add_page_break(); print("      Abstract Page added.")

    def _insert_placeholder(self, placeholder_text: str, heading_text: str, heading_style: str, doc_type: str):
        print(f"    Inserting Placeholder for: {heading_text}")
        self.add_formatted_paragraph(heading_text, heading_style, doc_type)
        p = self.add_formatted_paragraph(placeholder_text, "normal_text", doc_type)
        self.placeholder_paragraphs[placeholder_text] = p
        self.add_page_break()

    def insert_toc_placeholder(self, doc_type="report"): self._insert_placeholder(TOC_PLACEHOLDER, "Table of Contents", "heading_list_toc", doc_type)
    def insert_lof_placeholder(self, doc_type="report"): self._insert_placeholder(LOF_PLACEHOLDER, "List of Figures", "heading_list_toc", doc_type)
    def insert_lot_placeholder(self, doc_type="report"): self._insert_placeholder(LOT_PLACEHOLDER, "List of Tables", "heading_list_toc", doc_type)

    def add_heading(self, text: str, level: int, doc_type: str):
        if not text: return
        number_str, style_key, heading_text_final = "", "", text; is_numbered = False

        if doc_type == 'report':
            if level == 1:
                self.chapter_num += 1; self.section_num, self.subsection_num = 0, 0
                self.figure_num_in_chapter, self.table_num_in_chapter = 0, 0
                self.current_chapter_number = self.chapter_num
                style_key = "heading_chapter"; number_str = f"{self.chapter_num}"
                heading_text_final = f"CHAPTER {number_str}: {text.upper()}"; is_numbered = True
            elif level == 2 and self.chapter_num > 0:
                self.section_num += 1; self.subsection_num = 0
                style_key = "heading_section"; number_str = f"{self.chapter_num}.{self.section_num}"
                heading_text_final = f"{number_str} {text}"; is_numbered = True
            elif level == 3 and self.section_num > 0:
                self.subsection_num += 1
                style_key = "heading_subsection"; number_str = f"{self.chapter_num}.{self.section_num}.{self.subsection_num}"
                heading_text_final = f"{number_str} {text}"; is_numbered = True
            else:
                 print(f"Warning: Cannot correctly number heading L{level} ('{text}'). Adding unnumbered.")
                 style_key = "heading_subsection" if level >= 3 else ("heading_section" if level == 2 else "heading_chapter")
        elif doc_type == 'synopsis':
            if level == 1:
                self.chapter_num += 1; style_key = "heading1"
                number_str = f"{self.chapter_num}."; heading_text_final = f"{number_str} {text}"; is_numbered = True
            else: print(f"Warning: L{level} heading not standard for Synopsis ('{text}')."); style_key = "normal_text"

        print(f"    Adding Heading (L{level}, Num:{number_str or 'N/A'}, Style:{style_key}): {text}")
        p = self.add_formatted_paragraph(heading_text_final, style_key, doc_type)
        if p and is_numbered and style_key != "normal_text":
            self.headings.append({"level": level, "text": heading_text_final, "number": number_str, "paragraph": p})

    def add_figure(self, image_path_str: str, caption_text: str, doc_type="report"):
        print(f"    Adding Figure: {caption_text[:30]}...")
        if not caption_text: caption_text = "[No Caption Provided]"
        image_path = Path(image_path_str);
        if not image_path.is_absolute(): image_path = (Path.cwd() / image_path).resolve()
        if not image_path.is_file(): print(f"ERROR: Image file not found: {image_path}. Skipping."); return

        self.figure_num_in_chapter += 1
        figure_number_str = f"{self.current_chapter_number}.{self.figure_num_in_chapter}" if self.current_chapter_number > 0 else f"{self.figure_num_in_chapter}"
        figure_prefix = self.guideline_mgr.get_doc_rules(doc_type).get("figure_prefix", "Fig")
        full_caption = f"{figure_prefix} {figure_number_str}: {caption_text}"

        try:
            section = self.doc.sections[-1]
            available_width = section.page_width - section.left_margin - section.right_margin
            image_width = min(Inches(6.0), available_width * 0.95)
            self.doc.add_picture(str(image_path), width=image_width)
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e: print(f"ERROR: Could not add image {image_path}: {e}"); p = self.add_formatted_paragraph(f"[Figure Load Error]", "normal_text", doc_type)

        caption_paragraph = self.add_formatted_paragraph(full_caption, "caption", doc_type)
        if caption_paragraph: self.figures.append({"number": figure_number_str, "caption": caption_text, "full_caption": full_caption, "paragraph": caption_paragraph}); print(f"      Figure {figure_number_str} added and tracked.")

    def add_table(self, data: list, caption_text: str, doc_type="report", header=True):
        print(f"    Adding Table: {caption_text[:30]}...")
        if not caption_text: caption_text = "[No Caption Provided]"
        if not data or not isinstance(data, list) or not all(isinstance(row, list) for row in data): print("ERROR: Invalid table data. Skipping."); return

        self.table_num_in_chapter += 1
        table_number_str = f"{self.current_chapter_number}.{self.table_num_in_chapter}" if self.current_chapter_number > 0 else f"{self.table_num_in_chapter}"
        table_prefix = self.guideline_mgr.get_doc_rules(doc_type).get("table_prefix", "Table")
        full_caption = f"{table_prefix} {table_number_str}: {caption_text}"

        caption_paragraph = self.add_formatted_paragraph(full_caption, "caption", doc_type)
        if caption_paragraph: caption_paragraph.paragraph_format.space_after = Pt(4)

        try:
            num_rows, num_cols = len(data), len(data[0]) if data else 0
            if num_cols == 0: print("Warning: Table data has zero columns."); return
            table = self.doc.add_table(rows=num_rows, cols=num_cols); table.style = 'Table Grid'

            if header and num_rows > 0:
                for j, cell_text in enumerate(data[0]):
                    cell = table.cell(0, j); cell.text = str(cell_text)
                    if cell.paragraphs and cell.paragraphs[0].runs: cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            start_row = 1 if header else 0
            for i in range(start_row, num_rows):
                for j, cell_text in enumerate(data[i]): table.cell(i, j).text = str(cell_text)
            table.autofit = True
            self.add_formatted_paragraph("", "normal_text", doc_type).paragraph_format.space_before = Pt(12)
        except Exception as e: print(f"ERROR: Could not create table '{caption_text}': {e}"); import traceback; traceback.print_exc()

        if caption_paragraph: self.tables.append({"number": table_number_str, "caption": caption_text, "full_caption": full_caption, "paragraph": caption_paragraph}); print(f"      Table {table_number_str} added and tracked.")

    def _add_page_number_field(self, run, style='arabic'):
        """Helper to add PAGE field using Oxml with format switch."""
        fldChar_begin = OxmlElement('w:fldChar'); fldChar_begin.set(qn('w:fldCharType'), 'begin')
        fldChar_instr = OxmlElement('w:instrText')
        format_switch = ""
        # Reference: https://support.microsoft.com/en-us/office/field-codes-page-field-41 P-5b97-459a-ab31-65069f4c3983
        if style == 'roman_lower': format_switch = r'\* lowerroman' # Correct switch for lowercase Roman
        elif style == 'roman_upper': format_switch = r'\* upperroman' # Correct switch for uppercase Roman
        # Arabic (decimal) is the default if no format switch is provided
        fldChar_instr.text = f' PAGE {format_switch} ' # Simplified field code
        fldChar_sep = OxmlElement('w:fldChar'); fldChar_sep.set(qn('w:fldCharType'), 'separate')
        fldChar_end = OxmlElement('w:fldChar'); fldChar_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_begin); run._r.append(fldChar_instr); run._r.append(fldChar_sep); run._r.append(fldChar_end)

    def apply_page_numbering(self):
        print("    Applying Page Numbering...")
        if self.body_section_index == -1: print("Warning: Body section index not set. Applying Arabic to all."); self.body_section_index = 1
        num_rules = self.guideline_mgr.get_page_numbering_rules('report')
        front_format = num_rules.get('front_matter_format', 'roman_lower'); body_format = num_rules.get('body_format', 'arabic')
        position = num_rules.get('position', 'bottom_center')

        try:
            for i, section in enumerate(self.doc.sections):
                section.different_first_page_header_footer = False # Ensure consistency
                if i > 0: section.footer.is_linked_to_previous = True # Link footers
                footer = section.footer
                if not footer.paragraphs: footer.add_paragraph()
                footer_para = footer.paragraphs[0]; footer_para.clear()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # Assume bottom_center

                num_style = front_format if i < self.body_section_index else body_format
                run = footer_para.add_run()
                self._add_page_number_field(run, style=num_style) # Use the corrected field code helper

                num_font_style = self.guideline_mgr.get_formatting_rule('report', 'page_number')
                if num_font_style:
                    run.font.name = num_font_style.get('font', 'Times New Roman'); run.font.size = num_font_style.get('size', Pt(10))
                else: run.font.name = 'Times New Roman'; run.font.size = Pt(10)
                print(f"      Applied '{num_style}' page numbering to footer of Section {i}.")

            if self.body_section_index > 0 and self.body_section_index < len(self.doc.sections):
                body_section = self.doc.sections[self.body_section_index]
                # Check if pgNumType already exists, modify if it does, otherwise add
                sectPr = body_section._sectPr
                pgNumType = sectPr.find(qn('w:pgNumType'))
                if pgNumType is None:
                     pgNumType = OxmlElement('w:pgNumType')
                     sectPr.append(pgNumType)
                pgNumType.set(qn('w:start'), '1') # Set restart
                # Remove format attribute if it exists, let footer field handle format
                if qn('w:fmt') in pgNumType.attrib: del pgNumType.attrib[qn('w:fmt')]
                print(f"      Configured Section {self.body_section_index} to restart page numbering at 1.")

        except Exception as e: print(f"ERROR applying page numbering: {e}"); import traceback; traceback.print_exc()

    def _find_placeholder_paragraph(self, placeholder_text):
        if placeholder_text in self.placeholder_paragraphs: return self.placeholder_paragraphs[placeholder_text]
        for paragraph in self.doc.paragraphs:
            if placeholder_text == paragraph.text: self.placeholder_paragraphs[placeholder_text] = paragraph; return paragraph
        return None

    def _add_list_entry(self, text: str, indent_value: Inches, placeholder_para, item_style_key: str, doc_type: str):
        """Inserts a list entry before the placeholder with indentation and tab."""
        text_with_tab = f"{text}\t..." # Add placeholder dots
        new_para = placeholder_para.insert_paragraph_before(text_with_tab)
        self._apply_paragraph_format(new_para, item_style_key, doc_type)
        new_para.paragraph_format.left_indent = indent_value

        # Define right-aligned tab stop near the margin
        tab_stops = new_para.paragraph_format.tab_stops
        section = self.doc.sections[-1] # Use current section for margins
        page_width = section.page_width
        right_margin_pos = page_width - section.right_margin - Inches(0.1) # Slightly inside margin
        try:
            # Ensure position is valid (greater than left indent)
            if right_margin_pos > indent_value:
                 # Check if tab stop already exists, maybe clear first?
                 # tab_stops.clear_all() # Optional: clear existing tabs
                 tab_stops.add_tab_stop(right_margin_pos, WD_TAB_ALIGNMENT.RIGHT)
            else: print(f"      Warning: Cannot add tab stop for '{text[:30]}...' as position <= indent.")
        except Exception as ve: print(f"      Warning: Error adding tab stop for '{text[:30]}...': {ve}")

    def generate_toc(self, doc_type="report"):
        placeholder_para = self._find_placeholder_paragraph(TOC_PLACEHOLDER)
        if not placeholder_para: print(f"Warning: {TOC_PLACEHOLDER} not found."); return
        print(f"      Generating Table of Contents (including Levels 1-3)...")
        item_style_key = "list_entry"
        for heading_info in reversed(self.headings): # Reverse order for insert_before
             indent_level = heading_info.get('level', 1) - 1
             indent_value = Inches(0.4 * indent_level) # Adjust multiplier as needed
             text = heading_info.get('text', '[Missing Heading]')
             self._add_list_entry(text, indent_value, placeholder_para, item_style_key, doc_type)
        placeholder_para.text = ""; print(f"      TOC generation complete.") # Clear placeholder

    def generate_lof(self, doc_type="report"):
        placeholder_para = self._find_placeholder_paragraph(LOF_PLACEHOLDER)
        if not placeholder_para: print(f"Warning: {LOF_PLACEHOLDER} not found."); return
        print(f"      Generating List of Figures..."); item_style_key = "list_entry"
        for fig_info in reversed(self.figures):
             text = fig_info.get('full_caption', '[Missing Figure Caption]')
             self._add_list_entry(text, Inches(0), placeholder_para, item_style_key, doc_type) # No indent for LoF
        placeholder_para.text = ""; print(f"      LoF generation complete.")

    def generate_lot(self, doc_type="report"):
        placeholder_para = self._find_placeholder_paragraph(LOT_PLACEHOLDER)
        if not placeholder_para: print(f"Warning: {LOT_PLACEHOLDER} not found."); return
        print(f"      Generating List of Tables..."); item_style_key = "list_entry"
        for table_info in reversed(self.tables):
             text = table_info.get('full_caption', '[Missing Table Caption]')
             self._add_list_entry(text, Inches(0), placeholder_para, item_style_key, doc_type) # No indent for LoT
        placeholder_para.text = ""; print(f"      LoT generation complete.")

    def finalize_document(self):
        print("    Finalizing document: Generating Lists and applying Page Numbers...")
        self.generate_toc()
        self.generate_lof()
        self.generate_lot()
        self.apply_page_numbering()
        print("    Document finalized.")

    def save_document(self, filename: str):
        output_path = Path(filename); output_path.parent.mkdir(parents=True, exist_ok=True)
        try: self.doc.save(output_path); print(f"    Document successfully saved to: {output_path}")
        except PermissionError: print(f"ERROR: Permission denied saving to {output_path}. Is file open?")
        except Exception as e: print(f"ERROR: Failed to save document: {e}"); import traceback; traceback.print_exc()